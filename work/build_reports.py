from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('Deliverables')
TODAY = '2026-07-27'
NAVY = '17365D'
BLUE = '2F75B5'
LIGHT = 'EAF2F8'

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tc_pr.append(shd)

def set_cell(cell, text, header=False):
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.name = 'Aptos'; r.font.size = Pt(8.5)
            if header:
                r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255)
    if header: shade(cell, NAVY)

def setup(doc, title, subtitle):
    sec = doc.sections[0]
    sec.top_margin = Inches(.8); sec.bottom_margin = Inches(.7)
    sec.left_margin = Inches(.8); sec.right_margin = Inches(.8)
    normal = doc.styles['Normal']; normal.font.name='Aptos'; normal.font.size=Pt(10); normal.paragraph_format.space_after=Pt(6)
    for name,size,color in [('Title',24,NAVY),('Heading 1',16,NAVY),('Heading 2',12,BLUE),('Heading 3',10,BLUE)]:
        style=doc.styles[name]; style.font.name='Aptos Display' if name=='Title' else 'Aptos'; style.font.size=Pt(size); style.font.color.rgb=RGBColor.from_string(color)
    p=doc.add_paragraph(); p.style='Title'; p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.add_run(title)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(subtitle); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE); r.font.size=Pt(11)
    p=doc.add_paragraph(f'Version 1.0 | Prepared {TODAY} | Independent Verification & Validation'); p.runs[0].italic=True
    doc.add_paragraph('')

def bullet(doc, text): doc.add_paragraph(text, style='List Bullet')
def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for c,h in zip(t.rows[0].cells,headers): set_cell(c,h,True)
    for row in rows:
        cells=t.add_row().cells
        for c,v in zip(cells,row): set_cell(c,v)
    return t

def add_footer(doc):
    footer=doc.sections[0].footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Sitters4Me IV&V - Confidential working deliverable').font.size=Pt(8)

def review_report():
    d=Document(); setup(d,'Sitters4Me Requirements Review Report','Independent Verification & Validation (IV&V) - Requirements Baseline v1.0'); add_footer(d)
    d.add_heading('Executive summary',1)
    d.add_paragraph('The requirements specification provides a usable functional baseline for a two-sided babysitting marketplace. It identifies 180 uniquely numbered requirements across parent, sitter, dispatch, payments, notifications, administration, and security. The specification is strong on workflow detail but is not yet sufficient for a production readiness decision without measurable non-functional acceptance criteria, explicit failure handling, and a controlled execution environment.')
    d.add_heading('Assessment at a glance',1)
    table(d,['Area','Assessment','IV&V observation'],[
      ['Functional coverage','Good','Core marketplace paths and roles are described with requirement IDs.'],
      ['Testability','Moderate','Many outcomes are testable; timing, availability, and error criteria are often qualitative.'],
      ['Security specification','Moderate','Role isolation and PCI intent are stated; abuse prevention and session controls need detail.'],
      ['Reliability / recovery','Limited','Offline, dependency outage, retry, and reconciliation behavior is largely unspecified.'],
      ['Accessibility / privacy','Missing','No WCAG, retention, deletion, consent, or incident-response acceptance criteria.'],
    ])
    d.add_heading('Key gaps and recommendations',1)
    gaps=[
      ('Non-functional acceptance criteria','Define response-time targets for map load, dispatch, notifications, chat, and admin refresh; define supported concurrency, availability, and recovery objectives.'),
      ('Authentication and session policy','Specify password complexity, reset-token lifetime, rate limiting, lockout, session expiry, revocation, concurrent-session policy, and login audit events.'),
      ('Dispatch concurrency','Define atomic reservation behavior, duplicate-request prevention, cancellation-versus-accept race resolution, and idempotency keys for all lifecycle actions.'),
      ('Payment exception paths','Define authorization/capture timing, declined-card flow, retry ownership, duplicate-charge prevention, refunds, disputes, payout reconciliation, and webhook failure recovery.'),
      ('Location and safety','Define permission-denied behavior, accuracy threshold, stale-location handling, spoofing response, background location policy, and emergency escalation expectations.'),
      ('Data protection','Specify retention/deletion rules for account, child, payment, location, chat, and audit data; state encryption and access-log expectations.'),
      ('Accessibility','Adopt WCAG 2.2 AA targets, dynamic text/scaling acceptance criteria, screen-reader labels, focus order, contrast, and reduced-motion behavior.'),
      ('Operations','Document monitoring, alerting, maintenance mode, backups, restore objective, third-party outage communication, and incident ownership.'),
    ]
    table(d,['Gap','Recommendation'],gaps)
    d.add_heading('Implementation review observations',1)
    d.add_paragraph('Static inspection was performed on the supplied project snapshot. It is not a substitute for dynamic security testing or production execution.')
    table(d,['Reference','Observation','Risk / recommendation'],[
      ['Lint run','The project lint command reports 21 errors and 69 warnings.','Resolve blocking lint errors before a release candidate; track hook dependency warnings as reliability risks.'],
      ['app/index.tsx','The lint rule reports FileSystem.documentDirectory as unavailable from the imported namespace.','Validate against Expo SDK 54 file-system APIs; session restore is a critical requirement.'],
      ['Backend scope','Requirements reference auth.php, but it is not in the supplied source snapshot.','Obtain the deployed authentication service and configuration-free test environment for full API/security coverage.'],
      ['Source snapshot','Existing uncommitted modifications were present before IV&V work began.','Baseline the candidate build and record a commit/build identifier before execution.'],
    ])
    d.add_heading('Readiness conclusion',1)
    d.add_paragraph('Status: CONDITIONAL. The product can advance to controlled system testing after the lint blockers are resolved or risk-accepted, the backend/API baseline is supplied, and test accounts plus Stripe/Expo test configurations are available. A production go/no-go decision should wait for execution of the high-priority dispatch, payment, authorization, notification, and recovery suite.')
    d.add_heading('Assumptions and limitations',1)
    for x in ['This review is based on sitters4me_requirements.docx and the local source snapshot inspected on 2026-07-27.','No authenticated test environment, production logs, live device access, Stripe test credentials, database access, or push-notification evidence was supplied.','Execution status therefore remains Not Executed / Blocked where runtime access is required.'] : bullet(d,x)
    d.save(OUT/'01_Requirements_Review'/'Sitters4Me_IVV_Requirements_Review_Report_v1.0.docx')

def strategy():
    d=Document(); setup(d,'Sitters4Me Test Strategy and Test Plan','Independent Verification & Validation (IV&V) - v1.0'); add_footer(d)
    d.add_heading('Purpose and quality objectives',1)
    d.add_paragraph('This plan defines independent verification of the Sitters4Me mobile application, API-backed dispatch engine, Stripe payment integration, Expo notifications, and admin portal. The objective is to establish evidence that required behavior works, fails safely, preserves authorization boundaries, and recovers predictably.')
    d.add_heading('Scope',1)
    table(d,['In scope','Out of scope until access is supplied'],[
      ['Mobile parent and sitter workflows; admin portal; API contract behavior; session restoration; dispatch lifecycle; payments; notifications; data validation; role isolation; regression.','Production load certification; penetration testing beyond authorized test accounts; actual financial transfers; real end-user personal data; third-party service SLA certification.']
    ])
    d.add_heading('Test approach',1)
    for x in ['Requirements-based tests mapped one-to-one or one-to-many in the RTM.','Risk-based expansion for races, cancellation, third-party errors, payments, session recovery, authorization, and accessibility.','Static quality checks on the supplied code snapshot; dynamic mobile/API/admin execution when environment access is available.','Manual exploratory tests for maps, GPS, push notifications, usability, and cross-role workflows; automate stable API and smoke/regression paths.'] : bullet(d,x)
    d.add_heading('Priority model',1)
    table(d,['Priority','Definition','Examples'],[['P0','Safety, security, payment, or core booking blocker','Authorization, charge integrity, active-job state, dispatch acceptance'],['P1','Major feature failure with a workaround or bounded impact','Chat sync, scheduling, profile updates'],['P2','Secondary workflow or presentation defect','Non-blocking display, copy, preference behavior']])
    d.add_heading('Environment and test data needed',1)
    for x in ['Versioned mobile build for iOS and Android; web browser for admin portal.','Non-production API, database and Stripe test mode, with resettable seeded data.','At least two parent accounts, four sitter accounts (online/offline/pending/suspended), one admin, payment test cards, and accounts with/without active jobs.','Two physical devices or emulators capable of push notifications and location simulation.','Central log access or request correlation IDs for API, Stripe webhook, and Expo push investigations.'] : bullet(d,x)
    d.add_heading('Entry and exit criteria',1)
    table(d,['Entry','Exit'],[['Build identifier and deployment notes available; P0/P1 test data seeded; no known build-blocking installation issue; backend source/API contract available.','All P0 tests executed; P0 failures closed or formally accepted; ≥95% P1 executed; payment/dispatch/security evidence reviewed; open risks documented with owner and target date.']])
    d.add_heading('Defect management',1)
    d.add_paragraph('Defects are logged with reproducible steps, expected/actual behavior, environment, evidence, severity, priority, owner, and retest status. Severity is business impact; priority is remediation urgency. Any issue touching authorization, personal/child data, payment integrity, or incorrect active-job state is triaged as P0 until proven otherwise.')
    d.save(OUT/'03_Test_Strategy'/'Sitters4Me_Test_Strategy_and_Test_Plan_v1.0.docx')

def final_report():
    d=Document(); setup(d,'Sitters4Me IV&V Final Assessment','Current evidence-based assessment - v1.0'); add_footer(d)
    d.add_heading('Assessment statement',1)
    d.add_paragraph('This report consolidates the requirements review, traceability design, test design, static inspection, and execution readiness assessment completed against the supplied source snapshot. It is an interim final assessment: runtime execution remains pending because a deployed test environment, valid test credentials, and live device/service access were not supplied.')
    d.add_heading('Evidence completed',1)
    table(d,['Area','Status','Evidence'],[['Requirements review','Complete','180 uniquely identified requirements catalogued.'],['RTM and test design','Complete','Requirement-mapped baseline plus risk-based scenarios.'],['Static inspection','Complete','Lint run completed; 21 errors and 69 warnings observed.'],['Mobile/API/admin execution','Blocked','Requires test environment, credentials, test data, and device/service access.'],['Security / performance execution','Blocked','Requires explicitly authorized test target and monitoring instrumentation.']])
    d.add_heading('Release recommendation',1)
    d.add_paragraph('NOT READY FOR PRODUCTION APPROVAL. The recommendation is to proceed to a controlled system-test phase, not a public release, after resolving the blocking quality findings and providing the test environment. Reassess release readiness after all P0 tests pass and payment, dispatch, authorization, and notification end-to-end evidence is attached.')
    d.add_heading('Release gates',1)
    for x in ['Resolve or formally accept all lint errors; validate session persistence against Expo SDK 54.','Supply and test the complete authentication/API surface, including auth.php.','Execute P0 payment, dispatch race, cancellation, and authorization tests in Stripe test mode.','Verify push delivery and cold-start routing on physical iOS and Android devices.','Complete security tests only under explicit authorization and document results.','Establish monitoring, data retention/privacy controls, and incident/recovery runbooks.'] : bullet(d,x)
    d.add_heading('Outstanding risk',1)
    d.add_paragraph('Highest residual risks are payment duplication/failure recovery, dispatch and cancellation race conditions, authorization boundary enforcement, incorrect restoration of active jobs, stale or unavailable GPS data, and missing explicit resilience/accessibility requirements.')
    d.save(OUT/'08_Final_Report'/'Sitters4Me_IVV_Final_Assessment_v1.0.docx')

review_report(); strategy(); final_report()
