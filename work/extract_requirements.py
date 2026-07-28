from pathlib import Path
from docx import Document


source = Path("sitters4me_requirements.docx")
document = Document(source)
lines = []

for index, paragraph in enumerate(document.paragraphs, start=1):
    value = paragraph.text.strip()
    if value:
        style = paragraph.style.name if paragraph.style is not None else "(no style)"
        lines.append(f"P{index}\t{style}\t{value}")

for index, table in enumerate(document.tables, start=1):
    lines.append(f"\n[TABLE {index}]")
    for row in table.rows:
        lines.append(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))

Path("work/requirements_extracted.txt").write_text("\n".join(lines), encoding="utf-8")
print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} extracted={len(lines)}")
